import json
from collections import defaultdict

from docx import Document

from create_book import add_bookmark, add_hyperlink

with open("data.json") as f:
    subjects = {
        int(k): (
            "Economics" if "Economics" in " ".join(v["subject"]).split() else "Other"
        )
        for k, v in json.load(f).items()
    }


titles = {}
authors = {}
chunks_by_book = defaultdict(lambda: defaultdict(list))
with open("data_cleaned.json") as f:
    for book_data in json.load(f):
        titles[int(book_data["ebook_no"])] = book_data["title"]
        authors[int(book_data["ebook_no"])] = book_data["author"]
        for chunk in book_data["chunks"]:
            chunks_by_book[int(book_data["ebook_no"])][chunk["chunk_index"]].append(
                chunk["text"]
            )

chunk_texts = []
chunk_id_by_book = defaultdict(list)
book_ids = {}
for book_id, book_chunks in chunks_by_book.items():
    book_chunks = list(book_chunks.values())
    for chunk_i in range(len(book_chunks)):
        book_ids[len(chunk_texts)] = int(book_id)
        chunk_id_by_book[int(book_id)].append(len(chunk_texts))
        chunk_texts.append(
            "\n".join([chunk_text for chunk_text in book_chunks[chunk_i]])
        )


def get_chunk_info(chunk_id, divergence=None):
    info = {
        "chunk_id": chunk_id,
        "book_id": book_ids[chunk_id],
        "title": titles[book_ids[chunk_id]],
        "author": authors[book_ids[chunk_id]],
        #  "subject": subjects[book_ids[chunk_id]],
    }
    if divergence:
        info["divergence"] = divergence

    return info


def least_divergence(chunk_id):
    this_book_id = book_ids[chunk_id]
    this_author_name = authors[book_ids[chunk_id]]
    this_book = None
    this_author = None
    other = None
    with open(f"/mnt/e/js_scores/{chunk_id}.json") as f:
        for matched_id, divergence in sorted(json.load(f).items(), key=lambda x: x[1]):
            if book_ids[int(matched_id)] == this_book_id:
                if this_book is None:
                    this_book = (int(matched_id), divergence)
            elif authors[book_ids[int(matched_id)]] == this_author_name:
                if this_author is None:
                    this_author = (int(matched_id), divergence)
            elif not other:
                other = (int(matched_id), divergence)
            else:
                break
    return this_book, this_author, other


for book_id, chunk_ids in chunk_id_by_book.items():
    doc = Document()
    doc.add_heading(titles[book_id])
    for chunk_id in chunk_ids:
        # print(get_chunk_info(chunk_id, divergence=None))
        same_book, same_author, other = least_divergence(chunk_id)

        paragraph = doc.add_paragraph("")
        add_bookmark(paragraph=paragraph, bookmark_name=f"chunk{chunk_id}")

        # if same_book:
        #    print("book", get_chunk_info(*same_book))
        # if same_author:
        #    print("author", get_chunk_info(*same_author))
        if other:
            # print(chunk_id, other)
            chunk_data = get_chunk_info(*other)
            print(chunk_data)
            add_hyperlink(
                paragraph,
                f"{chunk_data['book_id']}.docx#chunk{chunk_data['chunk_id']}",
                f"Jensen-Shannon divergence: {chunk_data['divergence']}\n{chunk_data['title']} - {chunk_data['author']} \n#{chunk_data['chunk_id']}",
                chunk_texts[chunk_id],
            )
        else:
            paragraph.add_run(chunk_texts[chunk_id])
            # add_link(
            #    paragraph,
            #    f"0{chunk_data['book_id']}.docx#chunk{chunk_data['chunk_id']}",
            #    chunk_texts[chunk_id],
            #    tool_tip=f"Jensen-Shannon divergence: {chunk_data['divergence']}\n{chunk_data['title']} - {chunk_data['author']} \n#{chunk_data['chunk_id']}",
            # )
    doc.save(f"results/{book_id}.docx")
