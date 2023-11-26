from docx import Document

from docx import Document
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX


import json
from collections import defaultdict
import os


import os
import json
from collections import defaultdict, Counter

import pandas as pd
import numpy as np


def add_bookmark(paragraph, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
    start.set(docx.oxml.ns.qn("w:id"), "0")
    start.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(start)

    end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
    end.set(docx.oxml.ns.qn("w:id"), "0")
    end.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(end)


def add_link(paragraph, link_to, text, tool_tip=None):
    # create hyperlink node
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")

    # set attribute for link to bookmark
    hyperlink.set(
        docx.oxml.shared.qn("w:anchor"),
        link_to,
    )

    if tool_tip is not None:
        # set attribute for link to bookmark
        hyperlink.set(
            docx.oxml.shared.qn("w:tooltip"),
            tool_tip,
        )

    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True


titles = {}
authors = {}
chunks_by_book = defaultdict(lambda: defaultdict(list))
with open("data_cleaned.json") as f:
    for book_data in json.load(f):
        titles[int(book_data["ebook_no"])] = book_data["title"]
        authors[int(book_data["ebook_no"])] = book_data["author"]
        for chunk in book_data["chunks"]:
            chunks_by_book[int(book_data["ebook_no"])][chunk["chunk_index"]].append(
                chunk["text"]
            )

chunk_texts = []
chunk_id_by_book = defaultdict(list)
book_ids = {}
for book_id, book_chunks in chunks_by_book.items():
    book_chunks = list(book_chunks.values())
    for chunk_i in range(len(book_chunks)):
        book_ids[len(chunk_texts)] = int(book_id)
        chunk_texts.append(
            "\n".join([chunk_text for chunk_text in book_chunks[chunk_i]])
        )
        chunk_id_by_book[int(book_id)].append(chunk_i)


for book_id, chunk_ids in chunk_id_by_book.items():
    doc = Document()
    doc.add_heading(titles[book_id])

    for chunk_id in chunk_ids:
        paragraph = doc.add_paragraph("")
        add_bookmark(paragraph=paragraph, bookmark_name=f"chunk{chunk_id}")
        add_link(
            paragraph,
            f"chunk{chunk_id-1}",
            chunk_texts[chunk_id],
            tool_tip=chunk_texts[chunk_id],
        )

    doc.save(f"most_similar/{book_id}.docx")
