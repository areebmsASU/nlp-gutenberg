from docx import Document

import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

import os
import json
from collections import defaultdict, Counter

import pandas as pd
import numpy as np

doc = Document("most_similar/output.docx")

num_para = 0
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        print(run.text)
    # print(paragraph.style.name, paragraph.runs)
    # if not paragraph.style.name.startswith("Heading"):
    #    num_para += 1

print(num_para)

print(len(doc.paragraphs))

exit()


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    start = docx.oxml.shared.OxmlElement("w:bookmarkStart")
    start.set(docx.oxml.ns.qn("w:id"), "0")
    start.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement("w:bookmarkEnd")
    end.set(docx.oxml.ns.qn("w:id"), "0")
    end.set(docx.oxml.ns.qn("w:name"), bookmark_name)
    tag.append(end)


titles = {}
authors = {}
chunks_by_book = defaultdict(lambda: defaultdict(list))
with open("data_cleaned.json") as f:
    for book_data in json.load(f):
        titles[int(book_data["ebook_no"])] = book_data["title"]
        authors[int(book_data["ebook_no"])] = book_data["author"]
        for chunk in book_data["chunks"]:
            chunks_by_book[int(book_data["ebook_no"])][chunk["chunk_index"]].append(
                chunk["text"]
            )

chunk_texts = []
chunk_id_by_book = defaultdict(list)
book_ids = {}
for book_id, book_chunks in chunks_by_book.items():
    book_chunks = list(book_chunks.values())
    for chunk_i in range(len(book_chunks)):
        book_ids[len(chunk_texts)] = int(book_id)
        chunk_texts.append(
            "\n".join([chunk_text for chunk_text in book_chunks[chunk_i]])
        )
        chunk_id_by_book[int(book_id)].append(chunk_i)


doc = Document()


for book_id, chunk_ids in chunk_id_by_book.items():
    doc.add_heading(titles[book_id])

    for chunk_id in chunk_ids:
        paragraph = doc.add_paragraph("")
        add_bookmark(
            paragraph=paragraph,
            bookmark_text=f"{chunk_texts[chunk_id]}",
            bookmark_name=f"chunk{chunk_id}",
        )

    doc.add_page_break()


doc.save(r"most_similar/output.docx")
